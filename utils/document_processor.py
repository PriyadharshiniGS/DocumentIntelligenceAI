import os
import logging
import io
import PyPDF2
import docx
import pandas as pd
from PIL import Image
import pytesseract
import re

# Configure logging
logging.basicConfig(level=logging.DEBUG)
logger = logging.getLogger(__name__)

def process_document(file_path, file_name):
    """
    Process document based on file type and extract text content.
    
    Args:
        file_path (str): Path to the document file
        file_name (str): Original filename
    
    Returns:
        list: List of text chunks extracted from document
    """
    try:
        extension = file_name.split('.')[-1].lower()
        
        if extension == 'pdf':
            return process_pdf(file_path)
        elif extension == 'docx':
            return process_docx(file_path)
        elif extension == 'csv':
            return process_csv(file_path)
        else:
            raise ValueError(f"Unsupported document type: {extension}")
    
    except Exception as e:
        logger.error(f"Error processing document: {str(e)}")
        raise

def process_pdf(file_path):
    """
    Extract text from PDF file and split into chunks.
    
    Args:
        file_path (str): Path to the PDF file
    
    Returns:
        list: List of text chunks extracted from PDF
    """
    chunks = []
    try:
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            num_pages = len(pdf_reader.pages)
            
            # First pass: extract all text
            all_text = ""
            for page_num in range(num_pages):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    all_text += page_text + "\n\n"
            
            # Check if we got meaningful text
            if len(all_text.strip()) < 100 and num_pages > 0:
                logger.info("PDF text extraction yielded limited text. Trying OCR...")
                # If we didn't get much text, the PDF might be scanned images
                # Use Tesseract OCR with PDF2Image
                try:
                    import pdf2image
                    images = pdf2image.convert_from_path(file_path)
                    for i, image in enumerate(images):
                        page_text = pytesseract.image_to_string(image)
                        all_text += f"Page {i+1}:\n{page_text}\n\n"
                except ImportError:
                    logger.warning("pdf2image not available for OCR fallback")
            
            # Split into chunks of about 1000 words each
            chunks = split_text_into_chunks(all_text)
            
    except Exception as e:
        logger.error(f"Error processing PDF: {str(e)}")
        raise
    
    return chunks

def process_docx(file_path):
    """
    Extract text from DOCX file and split into chunks.
    
    Args:
        file_path (str): Path to the DOCX file
    
    Returns:
        list: List of text chunks extracted from DOCX
    """
    try:
        doc = docx.Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables_text = []
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' '.join([cell.text for cell in row.cells if cell.text.strip()])
                if row_text:
                    tables_text.append(row_text)
        
        # Combine and split into chunks
        all_text = "\n".join(paragraphs) + "\n" + "\n".join(tables_text)
        chunks = split_text_into_chunks(all_text)
        
        return chunks
    
    except Exception as e:
        logger.error(f"Error processing DOCX: {str(e)}")
        raise

def process_csv(file_path):
    """
    Extract content from CSV file and convert to text description.
    
    Args:
        file_path (str): Path to the CSV file
    
    Returns:
        list: List of text chunks describing CSV content
    """
    try:
        # Log the CSV processing attempt
        logger.debug(f"Attempting to process CSV file: {file_path}")
        
        # Read CSV in chunks to avoid memory issues
        chunks = []
        chunk_size = 100
        
        # First get column info and basic stats
        df_info = pd.read_csv(file_path, nrows=1)
        col_count = len(df_info.columns)
        overview = f"CSV file with {col_count} columns.\nColumn names: {', '.join(df_info.columns.tolist())}\n\n"
        chunks.append(overview)
        
        # Process the file in chunks
        for chunk_df in pd.read_csv(file_path, chunksize=chunk_size):
            chunk_text = []
            
            # Basic stats for numeric columns
            numeric_cols = chunk_df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                stats = chunk_df[numeric_cols].agg(['mean', 'min', 'max']).round(2)
                chunk_text.append("Numeric column statistics:")
                for col in numeric_cols:
                    chunk_text.append(f"{col}: mean={stats.loc['mean', col]}, min={stats.loc['min', col]}, max={stats.loc['max', col]}")
            
            # Sample rows (limited)
            sample_size = min(5, len(chunk_df))
            chunk_text.append("\nSample rows:")
            for _, row in chunk_df.head(sample_size).iterrows():
                row_items = []
                for col in chunk_df.columns:
                    val = row[col]
                    val_str = str(val) if pd.notna(val) else "NULL"
                    row_items.append(f"{col}: {val_str}")
                chunk_text.append(", ".join(row_items))
            
            chunks.append("\n".join(chunk_text))
        
        return chunks
        
        # Try to read the CSV with different encodings if needed
        try:
            df = pd.read_csv(file_path)
        except UnicodeDecodeError:
            logger.debug("UTF-8 encoding failed, trying latin1 encoding")
            df = pd.read_csv(file_path, encoding='latin1')
        
        # Get basic info about the CSV
        row_count = len(df)
        col_count = len(df.columns)
        logger.debug(f"Successfully read CSV with {row_count} rows and {col_count} columns")
        
        # Create a text description of the CSV data
        overview = f"CSV file with {row_count} rows and {col_count} columns.\n"
        overview += f"Column names: {', '.join(df.columns.tolist())}\n\n"
        
        # Add data sample (first few rows)
        sample_size = min(5, row_count)
        sample_rows = []
        
        for i in range(sample_size):
            row = df.iloc[i]
            row_items = []
            
            # Handle different data types properly
            for col in df.columns:
                try:
                    value = row[col]
                    # Convert numpy/pandas specific types to Python native types
                    if pd.isna(value):
                        formatted_value = "NULL"
                    elif isinstance(value, (int, float)):
                        # Format numeric values with appropriate precision
                        if float(value).is_integer():
                            formatted_value = str(int(value))
                        else:
                            formatted_value = f"{value:.4g}"
                    else:
                        # Convert to string and limit length
                        formatted_value = str(value)
                        if len(formatted_value) > 50:
                            formatted_value = formatted_value[:47] + "..."
                    
                    row_items.append(f"{col}: {formatted_value}")
                except Exception as item_err:
                    logger.warning(f"Error formatting column {col}: {str(item_err)}")
                    row_items.append(f"{col}: [ERROR: Could not format value]")
            
            sample_rows.append(f"Row {i+1}: {', '.join(row_items)}")
        
        overview += "Sample data:\n" + "\n".join(sample_rows) + "\n\n"
        
        # Add summary statistics for numeric columns (safer implementation)
        try:
            numeric_cols = df.select_dtypes(include=['number']).columns
            if len(numeric_cols) > 0:
                overview += "Summary statistics for numeric columns:\n"
                for col in numeric_cols:
                    try:
                        stats = df[col].describe()
                        # Format with safer string formatting that handles NaN values
                        stats_str = f"{col}: "
                        
                        for stat_name in ['min', 'max', 'mean', 'std']:
                            if stat_name in stats and not pd.isna(stats[stat_name]):
                                val = stats[stat_name]
                                if abs(val) < 0.01 or abs(val) > 1000:
                                    # Use scientific notation for very small or large numbers
                                    stats_str += f"{stat_name}={val:.2e}, "
                                else:
                                    stats_str += f"{stat_name}={val:.2f}, "
                            else:
                                stats_str += f"{stat_name}=N/A, "
                        
                        overview += stats_str.rstrip(', ') + "\n"
                    except Exception as stats_err:
                        logger.warning(f"Error calculating statistics for column {col}: {str(stats_err)}")
                        overview += f"{col}: [Error calculating statistics]\n"
        except Exception as stats_section_err:
            logger.warning(f"Error generating statistics section: {str(stats_section_err)}")
            overview += "Could not generate statistics for numeric columns due to an error.\n"
        
        # Process data in chunks if large
        chunks = [overview]
        logger.debug("Generated overview chunk for CSV")
        
        # If the CSV has many rows, process additional chunks (with enhanced safety)
        if row_count > 10:
            chunk_size = 20  # Reduced chunk size for safety
            for i in range(0, row_count, chunk_size):
                try:
                    end = min(i + chunk_size, row_count)
                    chunk_df = df.iloc[i:end]
                    
                    chunk_text = f"CSV data rows {i+1} to {end}:\n"
                    
                    # Process each row safely
                    for j, row in chunk_df.iterrows():
                        row_items = []
                        
                        for col in df.columns:
                            try:
                                value = row[col]
                                # Format values safely as we did above
                                if pd.isna(value):
                                    formatted_value = "NULL"
                                elif isinstance(value, (int, float)):
                                    if float(value).is_integer():
                                        formatted_value = str(int(value))
                                    else:
                                        formatted_value = f"{value:.4g}"
                                else:
                                    formatted_value = str(value)
                                    if len(formatted_value) > 50:
                                        formatted_value = formatted_value[:47] + "..."
                                
                                row_items.append(f"{col}: {formatted_value}")
                            except Exception:
                                row_items.append(f"{col}: [Error]")
                        
                        # Add this row to the chunk text
                        chunk_text += f"Row {j+1}: {', '.join(row_items)}\n"
                    
                    chunks.append(chunk_text)
                    logger.debug(f"Generated chunk for CSV rows {i+1}-{end}")
                    
                except Exception as chunk_err:
                    logger.warning(f"Error processing chunk {i//chunk_size + 1}: {str(chunk_err)}")
                    chunks.append(f"Error processing rows {i+1}-{min(i+chunk_size, row_count)}")
        
        logger.debug(f"CSV processing complete. Generated {len(chunks)} chunks")
        return chunks
    
    except Exception as e:
        logger.error(f"Error processing CSV: {str(e)}")
        import traceback
        logger.error(traceback.format_exc())
        # Return a simple chunk with the error information to avoid failing completely
        return [f"Error processing CSV file: {str(e)}"]

def split_text_into_chunks(text, words_per_chunk=1000):
    """
    Split text into chunks of approximately specified number of words.
    
    Args:
        text (str): Text to split
        words_per_chunk (int): Target number of words per chunk
    
    Returns:
        list: List of text chunks
    """
    # Clean the text - remove excessive whitespace and line breaks
    text = re.sub(r'\s+', ' ', text).strip()
    words = text.split()
    
    # If text is small enough, return as single chunk
    if len(words) <= words_per_chunk:
        return [text]
    
    chunks = []
    for i in range(0, len(words), words_per_chunk):
        chunk_words = words[i:i + words_per_chunk]
        chunks.append(' '.join(chunk_words))
    
    return chunks
